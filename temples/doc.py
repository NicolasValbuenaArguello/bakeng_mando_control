from docxtpl import DocxTemplate
from datetime import datetime
# from funciones.cargar_datos_tabla import *
import os
APP_PATH = os.getcwd()

directorio = APP_PATH +"/../datos/base_datos_no_borrar/temples/plantilla.docx"
doc = DocxTemplate(directorio)

def crear_word(date_fecha, linea_radiograma, linea_mando, com_ampliacion, date_fecha_hecho, time_hora_hecho, linea_unidad, linea_brigada, coordenadas, com_departamento, com_municipio, linea_lugar, linea_detalle_lugar, com_linea_operacion, linea_ordop, com_iniciativa, com_escalon_operacion, linea_cant_enemigo, com_denominacion, linea_estructura, com_operaciones_unificadas, com_tarea_TTF,com_tareas, com_tecnicas_actividades, com_amenaza, com_tipo_ubicacion,  com_informacion_hecho, linea_duracion_minutos, com_factores_de_inestabilidad, com_compania, linea_peloton, com_clase_soldados, com_personal_comprometido, linea_comandante, linea_telefono, com_que_produjo_exito, checkBox_principal, checkBox_complementaria, checkBox_vigilancia, checkBox_reconocimiento, comboBox_disciplina_inteligencia, com_tarea_inteligencia, textEdit_resumen, textEdit_gasto_municion, checkBox_narcotrafico, checkBox_groi, checkBox_policia_judicial, checkBox_Asalto_aereo, checkBox_apoyo_exde, checkBox_erradeicacion, checkBox_APOYO_BRCOM, accion_davaa, accion_conat, accion_ccoes, apoyo_divfe, apoyo_art, apoyo_bafur, apoyo_blica, inteligencia_militar, marco_conjunta, marco_interagencial, marco_multinacional, linea_davaa, linea_divfe, linea_ccoes, linea_conat, linea_art, linea_bafur, linea_blica, linea_inteligencia, linea_apoyo_aereo, linea_spoa, apoyo_aereo, combo_marco_conjunta_fuerza, combo_marco_interagencial_organizacion, combo_marco_multinacional_pais, combo_marco_multinacional_fuerza, combo_marco_multinacional_organizacion, linea_marco_multinacional_unidad, linea_marco_interagencial_unidad, linea_marco_conjunta_unidad, tabla, fila, tabla_1, fila_1, cdte, ofi_operaciones, sub_operaciones,tipo_resultados, apoyo_coeej, apoyo_fudat,ambiente_operacional, tipo_terreno, acta_incautacion, num):

    nombre = linea_radiograma + " - " + com_informacion_hecho
    continiuacion_hr = linea_radiograma + "/"+ linea_mando

    hola = ( 
        (1, 'Geek 1')

    ) 

    context = { 
            'date_fecha':date_fecha,
            'linea_radiograma':linea_radiograma,
            'linea_mando' : linea_mando,
            'com_ampliacion':com_ampliacion,
            'date_fecha_hecho':date_fecha_hecho,
            'time_hora_hecho':time_hora_hecho,
            'linea_unidad':linea_brigada ,
            'linea_brigada':linea_unidad,
            'coordenadas':coordenadas,
            'com_departamento':com_departamento,
            'com_municipio':com_municipio,
            'linea_lugar':linea_lugar,
            'linea_detalle_lugar':linea_detalle_lugar,
            'com_linea_operacion':com_linea_operacion,
            'linea_ordop':linea_ordop,
            'com_iniciativa':com_iniciativa,
            'com_escalon_operacion':com_escalon_operacion,
            'linea_cant_enemigo':linea_cant_enemigo,
            'com_denominacion':com_denominacion,
            'linea_estructura':linea_estructura,
            'com_operaciones_unificadas':com_operaciones_unificadas,
            'com_tarea_TTF':com_tarea_TTF,
            'com_tareas':com_tareas,
            'com_tecnicas_actividades':com_tecnicas_actividades,
            'com_amenaza':com_amenaza,
            'com_tipo_ubicacion':com_tipo_ubicacion,
            'com_informacion_hecho':com_informacion_hecho,
            'linea_duracion_minutos': linea_duracion_minutos,
            'com_factores_de_inestabilidad':com_factores_de_inestabilidad,
            'com_compania':com_compania,
            'linea_peloton':linea_peloton,
            'com_clase_soldados':com_clase_soldados,
            'com_personal_comprometido':com_personal_comprometido,
            'linea_comandante':linea_comandante,
            'linea_telefono':linea_telefono,
            'com_que_produjo_exito':com_que_produjo_exito,
            'checkBox_principal':checkBox_principal,
            'checkBox_complementaria':checkBox_complementaria,
            'checkBox_vigilancia':checkBox_vigilancia,
            'checkBox_reconocimiento':checkBox_reconocimiento,
            'comboBox_disciplina_inteligencia':comboBox_disciplina_inteligencia,
            'comboBox_disciplina_inteligencia':comboBox_disciplina_inteligencia,
            'com_tarea_inteligencia':com_tarea_inteligencia,
            'textEdit_resumen':textEdit_resumen,
            'textEdit_gasto_municion':textEdit_gasto_municion,
            'continiuacion_hr':continiuacion_hr,
            'checkBox_narcotrafico':checkBox_narcotrafico,
            'checkBox_groi':checkBox_groi,
            'checkBox_policia_judicial':checkBox_policia_judicial,
            'checkBox_Asalto_aereo':checkBox_Asalto_aereo,
            'checkBox_apoyo_exde':checkBox_apoyo_exde,
            'checkBox_erradeicacion':checkBox_erradeicacion,
            'checkBox_APOYO_BRCOM':checkBox_APOYO_BRCOM,
            'apoyo_fudat':apoyo_fudat,
            'accion_davaa':accion_davaa,
            'linea_davaa':linea_davaa,
            'accion_conat':accion_conat,
            'linea_conat':linea_conat,
            'accion_ccoes':accion_ccoes,
            'linea_ccoes':linea_ccoes,
            'apoyo_divfe':apoyo_divfe,
            'apoyo_divfe':apoyo_divfe,
            'linea_divfe':linea_divfe,
            'apoyo_art':apoyo_art,
            'linea_art':linea_art,
            'apoyo_aereo':apoyo_aereo,
            'linea_apoyo_aereo':linea_apoyo_aereo,
            'apoyo_bafur':apoyo_bafur,
            'linea_bafur':linea_bafur,
            'inteligencia_militar':inteligencia_militar,
            'linea_inteligencia':linea_inteligencia,
            'apoyo_blica':apoyo_blica,
            'linea_blica':linea_blica,
            'apoyo_coeej':apoyo_coeej,
            'marco_conjunta':marco_conjunta,
            'combo_marco_conjunta_fuerza':combo_marco_conjunta_fuerza,
            'linea_marco_conjunta_unidad':linea_marco_conjunta_unidad,
            'marco_interagencial':marco_interagencial,
            'combo_marco_interagencial_organizacion':combo_marco_interagencial_organizacion,
            'linea_marco_interagencial_unidad':linea_marco_interagencial_unidad,
            'marco_multinacional':marco_multinacional,
            'combo_marco_multinacional_pais':combo_marco_multinacional_pais,
            'combo_marco_multinacional_fuerza':combo_marco_multinacional_fuerza,
            'combo_marco_multinacional_organizacion':combo_marco_multinacional_organizacion,
            'linea_marco_multinacional_unidad':linea_marco_multinacional_unidad,
            'tipo_resultados':tipo_resultados,
            'linea_spoa':linea_spoa,
            'acta_incautacion':acta_incautacion,
            'ambiente_operacional':ambiente_operacional,
            'tipo_terreno':tipo_terreno,

            'hola' : hola[1]
                
    }
    doc.render(context)

    import docx 
    
    doc2 = docx.Document() 
    
    doc2.add_heading('GeeksForGeeks', 0) 
    

    
    # table = doc.add_table(rows=1, cols=2) 
    # table.style = 'resultados' 
    # row = table.rows[0].cells 
    # row[0].text = 'Id'
    # row[1].text = 'Name'
    
    # for id, name in tabla_1: 
    
        
    #     row = table.add_row().cells 
        
    #     row[0].text = str(id) 
    #     row[1].text = name 

    # # Añadimos un párrafo
    # parrafo = doc.add_paragraph()
    # # parrafo.add_run().add_break()
    if tabla_1 != []:
        fila_1 +=1
        for j in range(fila_1):
            table2 = doc.add_table(rows=1, cols=20) 
            table2.style = 'resultados' 
            row = table2.rows[0].cells 
            # row[0].text = 'Id'
            # row[1].text = 'Name'
            # row = table2.add_row().cells 

            row[0].paragraphs[0].add_run("Acción").bold = True
            row[6].merge(row[1])
            row[1].text = str(tabla_1[j][3])
            row[7].paragraphs[0].add_run("Tipo").bold = True
            row[9].merge(row[8])
            row[8].text = str(tabla_1[j][0])
            row[10].paragraphs[0].add_run("Subtipo").bold = True
            row[12].merge(row[11])
            row[11].text = str(tabla_1[j][1])
            row[13].paragraphs[0].add_run("Clase").bold = True
            row[16].merge(row[14])
            row[14].text = str(tabla_1[j][2])
            row[18].merge(row[17])
            row[17].paragraphs[0].add_run("Judicializado").bold = True
            row[19].text = str(tabla_1[j][4])

            row = table2.add_row().cells 

            row[0].paragraphs[0].add_run("Grado").bold = True 
            row[1].text = str(tabla_1[j][5])
            row[2].paragraphs[0].add_run("Nombres").bold = True  
            row[6].merge(row[3])
            row[3].text = str(tabla_1[j][6])
            row[7].paragraphs[0].add_run("Apellidos").bold = True   
            row[11].merge(row[8])
            row[8].text = str(tabla_1[j][7])
            row[13].merge(row[12])
            row[12].paragraphs[0].add_run("T. Documento").bold = True
            row[14].text = str(tabla_1[j][8])
            row[16].merge(row[15])
            row[15].text = str(tabla_1[j][9])
            row[17].paragraphs[0].add_run("Alias").bold = True
            row[19].merge(row[18])
            row[18].text = str(tabla_1[j][10])

            row = table2.add_row().cells 

            row[0].paragraphs[0].add_run("Sexo").bold = True
            row[2].merge(row[1])
            row[1].text = str(tabla_1[j][11])
            row[4].merge(row[3])
            row[3].paragraphs[0].add_run("Estado Civil").bold = True 
            row[6].merge(row[5])
            row[5].text = str(tabla_1[j][12])
            row[7].paragraphs[0].add_run("Edad").bold = True 
            row[9].merge(row[8])
            row[8].text = str(tabla_1[j][13])
            row[11].merge(row[10])
            row[10].paragraphs[0].add_run("Escolaridad").bold = True 
            row[14].merge(row[12])
            row[12].text = str(tabla_1[j][14])
            row[16].merge(row[15])
            row[15].paragraphs[0].add_run("No SIGAHD").bold = True 
            row[19].merge(row[17])
            row[17].text = str(tabla_1[j][15])

            row = table2.add_row().cells 

            row[1].merge(row[0])
            row[0].paragraphs[0].add_run("Motivo Entrega").bold = True 
            row[5].merge(row[2])
            row[2].text = str(tabla_1[j][16])
            row[7].merge(row[6])
            row[6].paragraphs[0].add_run("Portaba Arma").bold = True 
            row[8].text = str(tabla_1[j][17])
            row[10].merge(row[9])
            row[9].paragraphs[0].add_run("Tipo Arma").bold = True 
            row[13].merge(row[11])
            row[11].text = str(tabla_1[j][18])
            row[14].paragraphs[0].add_run("Herido").bold = True
            row[15].text = str(tabla_1[j][19])
            row[17].merge(row[16])
            row[16].paragraphs[0].add_run("Especialidad").bold = True
            row[19].merge(row[18])
            row[18].text = str(tabla_1[j][20])

            row = table2.add_row().cells 

            row[1].merge(row[0])
            row[0].paragraphs[0].add_run("Nivel Jerárquico").bold = True
            row[6].merge(row[2])
            row[2].text = str(tabla_1[j][21])
            row[9].merge(row[7])
            row[7].paragraphs[0].add_run("Tiempo Organización").bold = True
            row[12].merge(row[10])
            row[10].text = str(tabla_1[j][22])
            row[15].merge(row[13])
            row[13].paragraphs[0].add_run("Autoridad Disposición").bold = True
            row[19].merge(row[16])
            row[16].text = str(tabla_1[j][23])

            row = table2.add_row().cells 

            row[1].merge(row[0])
            row[0].paragraphs[0].add_run("Lugar Nacimiento").bold = True
            row[9].merge(row[2])
            row[2].text = str(tabla_1[j][24])
            row[11].merge(row[10])
            row[10].paragraphs[0].add_run("Delito Uno").bold = True
            row[19].merge(row[12])
            row[12].text = str(tabla_1[j][25])

            row = table2.add_row().cells 
            
            row[1].merge(row[0])
            row[0].paragraphs[0].add_run("Delito Dos").bold = True
            row[9].merge(row[2])
            row[2].text = str(tabla_1[j][26])
            row[11].merge(row[10])
            row[10].paragraphs[0].add_run("Delito Tres").bold = True
            row[19].merge(row[12])
            row[12].text = str(tabla_1[j][27])
            
            row = table2.add_row().cells 
            
            row[1].merge(row[0])
            row[0].text = str("Delito Cuatro")
            row[0].paragraphs[0].add_run("Delito Cuatro").bold = True
            row[9].merge(row[2])
            row[2].text = str(tabla_1[j][28])
            row[11].merge(row[10])
            row[10].paragraphs[0].add_run("Delito Cinco").bold = True
            row[19].merge(row[12])
            row[12].text = str(tabla_1[j][29])

            row = table2.add_row().cells 
      
            row[1].merge(row[0])
            row[0].paragraphs[0].add_run("Observación").bold = True
            row[19].merge(row[2])
            row[2].text = str(tabla_1[j][30])

            parrafo = doc.add_paragraph()


    if tabla != []:
        fila +=1
        for i in range(fila):
            table3 = doc.add_table(rows=1, cols=20) 
            table3.style = 'resultados' 
            row = table3.rows[0].cells 

            row[0].paragraphs[0].add_run("Acción").bold = True
            row[5].merge(row[1])
            row[1].text = str(tabla[i][3])
            row[6].paragraphs[0].add_run("Tipo").bold = True 
            row[10].merge(row[7])
            row[7].text = str(tabla[i][0])
            row[11].paragraphs[0].add_run("Subtipo").bold = True  
            row[15].merge(row[12])
            row[12].text = str(tabla[i][1])
            row[16].paragraphs[0].add_run("Clase").bold = True  
            row[19].merge(row[17])
            row[17].text = str(tabla[i][2])

            row = table3.add_row().cells 
            row[0].paragraphs[0].add_run("Cantidad").bold = True  
            row[2].merge(row[1])
            row[1].text = str(tabla[i][4])
            row[3].text = str(tabla[i][5])
            row[4].paragraphs[0].add_run("Calibre").bold = True  
            row[6].merge(row[5])
            row[5].text = str(tabla[i][6])
            row[8].merge(row[7])
            row[7].paragraphs[0].add_run("No. arma").bold = True  
            row[11].merge(row[9])
            row[9].text = str(tabla[i][7])
            row[12].paragraphs[0].add_run("Marca").bold = True  
            row[15].merge(row[13])
            row[13].text = str(tabla[i][8])
            row[17].merge(row[16])
            row[16].paragraphs[0].add_run("Fabricación").bold = True  
            row[19].merge(row[18])
            row[18].text = str(tabla[i][9])

            row = table3.add_row().cells 

            row[0].paragraphs[0].add_run("Valor").bold = True
            row[3].merge(row[1])
            row[1].text = str(tabla[i][10])

            row[5].merge(row[4])
            # # row[6].text = str("Observación")
            row[4].paragraphs[0].add_run("Observación").bold = True  
            row[19].merge(row[6])
            row[6].text = str(tabla[i][11])
            

            parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()

    table3 = doc.add_table(rows=1, cols=15) 
    table3.style = 'firmas' 
    row = table3.rows[0].cells 

    row[4].merge(row[0])
    row[0].paragraphs[0].add_run(sub_operaciones).bold = True
    row[9].merge(row[5])
    row[5].paragraphs[0].add_run(ofi_operaciones).bold = True
    row[14].merge(row[10])
    row[10].paragraphs[0].add_run(cdte).bold = True

    row = table3.add_row().cells 
    row[4].merge(row[0])
    row[0].paragraphs[0].add_run("Suboficial de Operaciones " + linea_brigada).bold = False
    # row[0].text = str()
    row[9].merge(row[5])
    row[5].paragraphs[0].add_run("Oficial de Operaciones  "+ linea_brigada).bold = False
    # row[5].text = str("Oficial de Operaciones  "+ linea_unidad)
    row[14].merge(row[10])
    row[10].paragraphs[0].add_run("Comandante " + linea_brigada).bold = False
    # row[10].text = str("Comandante " + linea_unidad)


    directorio = APP_PATH+"/../datos/WORD"
    fecha = datetime.today()

    fecha = fecha.strftime('%d-%m-%Y   %H-%M-%S')
    if num == 1:
        doc.save(directorio+"/" + nombre +" " +fecha +".docx")
    elif num == 2:
        doc.save(directorio+"/" + nombre + " " +fecha  +"(editado).docx")
from docxtpl import DocxTemplate
from datetime import datetime
try:
    from base_datos.base_datos import *
except ModuleNotFoundError:
    pass
#from mapa.distancias import *
import os
APP_PATH = os.getcwd()

directorio = APP_PATH +"/../datos/base_datos_no_borrar/temples/plantilla_maniobra.docx"
doc = DocxTemplate(directorio)

def crear_word_informe_inistop(self, numero_registro_informe, fecha_informe, linea_mando, cdte, ofi_operaciones, sub_operaciones, linea_unidad):
                    

    fecha=datetime.strptime(fecha_informe, '%d/%m/%Y')
    fecha = fecha.strftime('%d-%m-%Y ')
    nombre = "Informe de maniobra "+ " - " + str(fecha) +" - " +str(numero_registro_informe) 
    continiuacion_hr = str(numero_registro_informe) + "/"+ str(linea_mando)

    hola = ( 
        (1, 'Geek 1')

    ) 

    context = { 
            'fecha_elaboracion':fecha_informe,
            'registro_unidad':numero_registro_informe,
            'línea_unidad' : linea_mando,
            'unidad':linea_unidad,
            'continuación_doc' : continiuacion_hr,

            'hola' : hola[1]
                
    }
    doc.render(context)

    import docx 
    
    doc2 = docx.Document() 
    
    doc2.add_heading('GeeksForGeeks', 0) 
    

    
    # table = doc.add_table(rows=1, cols=2) 
    # table.style = 'resultados' 
    # row = table.rows[0].cells 
    # row[0].text = 'Id'
    # row[1].text = 'Name'
    
    # for id, name in tabla_1: 
    
        
    #     row = table.add_row().cells 
        
    #     row[0].text = str(id) 
    #     row[1].text = name 

    # # Añadimos un párrafo
    # parrafo = doc.add_paragraph()
    # # parrafo.add_run().add_break()

    date_fecha_final=datetime.strptime(fecha_informe, '%d/%m/%Y')

    date_fecha_inicial = date_fecha_final - timedelta(days=1)

    

    unidad_pel =  INSITOP_FECHA_UNIDAD(date_fecha_inicial, date_fecha_final)
    informacion_distacia =[]
    distacia =[]
    for unidad_cp in unidad_pel:
            unidades = INSITOP_FECHA_UNIDAD_FECHA(date_fecha_final, date_fecha_final,unidad_cp[0], unidad_cp[1] )


            for dato in unidades:

                unida_a = INSITOP_FECHA_UNIDAD_FECHA(date_fecha_final, date_fecha_final,unidad_cp[0], unidad_cp[1])

                unida_b = INSITOP_FECHA_UNIDAD_FECHA(date_fecha_inicial, date_fecha_inicial,unidad_cp[0], unidad_cp[1])

                exde = dato[18]
                cdte_cp = dato[19]
                tel = dato[20]
                efectivos =  str(dato[13]) + " - " + str(dato[14]) + " - " + str(dato[15]) + " - " + str(dato[16])+ " - " + str(dato[17])
                ordop = dato[27]
                departamento = dato[12]
                municipio  = dato[11]
                lugar  = dato[10]
                km= 0
                cordenada_n_act = 0
                cordenada_n_ant = 0

                x1=0
                x2=0
                y1=0
                y2=0

                
                if unida_a   :
                    for x in unida_a:
                        cordenadas_1 = coordenadas(self, x)
                        numero_coor = float(x[4])
                        ln =""
                        if numero_coor < 0:
                            ln  = "LS"
                        else:
                            ln  = "LN"

                        cordenada_n_act = str(ln) +" "+ str(x[4])+"° "+str(x[5])+"' "+str(x[6]) +"'' - "+str("LW") +" "+ str(x[7])+"° "+str(x[8])+"' "+str(x[9])+"''"
                        
                if  unida_b :      
                    for y in unida_b:
                        cordenadas_2 = coordenadas(self, y)

                        numero_coor = float(y[4])
                        ln =""
                        if numero_coor < 0:
                            ln  = "LS"
                        else:
                            ln  = "LN"

                        cordenada_n_ant = str(ln) +" "+ str(y[4])+"° "+str(y[5])+"' "+str(y[6]) +"'' - "+str("LW") +" "+ str(y[7])+"° "+str(y[8])+"' "+str(y[9])+"''"

                    x1=cordenadas_1[0]
                    y1=cordenadas_1[1]
                    x2=cordenadas_2[0]
                    y2=cordenadas_2[1]
                            
                    punto_1 = (x1, y1)
                    punto_2 = (x2, y2)

                    km =  distacia_puntos(punto_1, punto_2)

            fecha_informe_f=datetime.strptime(fecha_informe, '%d/%m/%Y')
            fecha_informe_f = fecha_informe_f.strftime('%d-%m-%Y ')

            date_fecha_inicial_f = date_fecha_inicial.strftime('%d-%m-%Y ')
                    

            json = (dato[1], unidad_cp[0], km, str(fecha_informe_f),str(date_fecha_inicial_f), cordenada_n_act, cordenada_n_ant, unidad_cp[1], efectivos, exde, cdte_cp, tel, ordop, departamento, municipio, lugar)
            informacion_distacia.append(json)

                # distacia.append(informacion_distacia)

    for dato in informacion_distacia:

        table2 = doc.add_table(rows=1, cols=30) 
        table2.style = 'unidades' 
        row = table2.rows[0].cells 
        
        row[3].merge(row[0])
        row[0].paragraphs[0].add_run("FECHA INICIAL").bold = True
        row[6].merge(row[4])
        row[4].text = str(dato[4])
                
        row[9].merge(row[7])
        row[7].paragraphs[0].add_run("FECHA FINAL").bold = True
        row[12].merge(row[10])
        row[10].text = str(dato[3])

        row[15].merge(row[13])
        row[13].paragraphs[0].add_run("COMPAÑIA").bold = True
        row[16].merge(row[16])
        row[16].text = str(dato[1])

        row[19].merge(row[17])
        row[17].paragraphs[0].add_run("PELOTÓN").bold = True
        row[20].merge(row[20])
        row[20].text = str(dato[7])
        
        row[23].merge(row[21])
        row[21].paragraphs[0].add_run("EFECTIVOS").bold = True
        row[26].merge(row[24])
        row[24].text = str(dato[8])
                
        row[28].merge(row[27])
        row[27].paragraphs[0].add_run("EXDE").bold = True
        row[29].merge(row[29])
        row[29].text = str(dato[9])

        row = table2.add_row().cells 

        row[1].merge(row[0])
        row[0].paragraphs[0].add_run("CDTE").bold = True
        row[10].merge(row[2])
        row[2].text = str(dato[10])

        
        row[13].merge(row[11])
        row[11].paragraphs[0].add_run("TELÉFONO").bold = True
        row[17].merge(row[14])
        row[14].text = str(dato[11])

    
        row[23].merge(row[18])
        row[18].paragraphs[0].add_run("COORDENADAS INICIALES").bold = True
        row[29].merge(row[24])
        row[24].text = str(dato[6])

        row = table2.add_row().cells 
   
        row[5].merge(row[0])
        row[0].paragraphs[0].add_run("COORDENADAS FINALES").bold = True
        row[11].merge(row[6])
        row[6].text = str(dato[5])

        row[14].merge(row[12])
        row[12].paragraphs[0].add_run("DISTANCIA").bold = True
        row[19].merge(row[15])
        km = format(dato[2], '0.2f')
        km = str(km)+" metros"
        row[15].text = str(km)
                   
        row[22].merge(row[20])
        row[20].paragraphs[0].add_run("ORDOP").bold = True
        row[29].merge(row[23])
        km = format(dato[2], '0.2f')
        row[23].text = str(dato[12])


        row = table2.add_row().cells 
           
        row[4].merge(row[0])
        row[0].paragraphs[0].add_run("DEPARTAMENTO").bold = True
        row[13].merge(row[5])
        row[5].text = str(dato[13])

                   
        row[16].merge(row[14])
        row[14].paragraphs[0].add_run("MUNICIPIO").bold = True
        row[29].merge(row[17])
        row[17].text = str(dato[14])

        row = table2.add_row().cells 

        row[2].merge(row[0])
        row[0].paragraphs[0].add_run("LUGAR").bold = True
        row[6].merge(row[3])
        row[3].text = str(dato[15])

        row[9].merge(row[7])
        row[7].paragraphs[0].add_run("OBSERVACIÓN").bold = True
        row[29].merge(row[10])

        parrafo = doc.add_paragraph() 

          

    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()

    table3 = doc.add_table(rows=1, cols=15) 
    table3.style = 'firmas' 
    row = table3.rows[0].cells 

    row[4].merge(row[0])
    row[0].paragraphs[0].add_run(sub_operaciones).bold = True
    row[9].merge(row[5])
    row[5].paragraphs[0].add_run(ofi_operaciones).bold = True
    row[14].merge(row[10])
    row[10].paragraphs[0].add_run(cdte).bold = True

    row = table3.add_row().cells 
    row[4].merge(row[0])
    row[0].paragraphs[0].add_run("Suboficial de Operaciones " + linea_unidad).bold = False
    # row[0].text = str()
    row[9].merge(row[5])
    row[5].paragraphs[0].add_run("Oficial de Operaciones  "+ linea_unidad).bold = False
    # row[5].text = str("Oficial de Operaciones  "+ linea_unidad)
    row[14].merge(row[10])
    row[10].paragraphs[0].add_run("Comandante " + linea_unidad).bold = False
    # row[10].text = str("Comandante " + linea_unidad)


    directorio = APP_PATH+"/../datos/WORD"

    doc.save(directorio+"/" + nombre +".docx")

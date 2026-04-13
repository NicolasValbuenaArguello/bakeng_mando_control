from docxtpl import DocxTemplate
from datetime import datetime
from base_datos.base_datos import *
#from mapa.distancias import *
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
import os
APP_PATH = os.getcwd()
cursor = b_d_i

concion_bd = Base_datos_central
def query_bd(self):
    def comando_query(self, **kwargs):
        datos = concion_bd.read_query_dinamico(self, **kwargs)
        return datos
    return  comando_query


def query_bd_dinamico(self):
    def comando_query(self, **kwargs):
        datos = concion_bd.select_dinamico(self, **kwargs)
        return datos
    return  comando_query

def query_bd_personal(self):
    def comando_query(self, **kwargs):
        datos = cursor.read_query_dinamico(self, **kwargs)
        return datos
    return  comando_query

#Llenado de los datos del primer formulario
@query_bd
def datos_base_datos(self, **kwargs):
    pass


@query_bd_dinamico
def select_dinamico_bd(self, **kwargs):
    pass
@query_bd_personal
def select_dinamico_bd_personal(self, **kwargs):
    pass

def numero_diferente_cero(x):
    if x !='':
        return x
    else:
        return 0

directorio = APP_PATH +"/../datos/base_datos_no_borrar/temples/parte_personal.docx"
doc = DocxTemplate(directorio)
def suma(i):
      
            TC =0
            MY = 0
            CT = 0
            TE = 0
            ST = 0
            SP = 0
            SV = 0
            SS = 0
            CP = 0
            CS = 0
            C3 = 0

            if i[52]=="TC":
                    TC = TC +1 
            elif i[52]=="MY":
                    MY = MY +1
            elif i[52]=="CT":
                    CT = CT +1
            elif i[52]=="TE":
                    TE = TE +1
            elif i[52]=="ST":
                    ST = ST +1
            elif i[52]=="SP":
                    SP = SP +1 
            elif i[52]=="SV":
                    SV = SV +1
            elif i[52]=="SS":
                    SS = SS +1
            elif i[52]=="CP":
                    CP = CP +1
            elif i[52]=="CS":
                    CS = CS +1 
            elif i[52]=="C3":
                    C3 = C3 +1 
            sumas = TC + MY + CT + TE + ST + SP + SV + SS + CP + CS + C3
            return[sumas, TC , MY , CT , TE , ST , SP , SV , SS , CP , CS , C3]



def crear_word_parte_personal(self, numero_registro_informe, fecha_informe, linea_mando, cdte, ofi_operaciones, sub_operaciones, divisiones, brigada, batallon, BATALLON_RELACION, FCG):
                    

    fecha=datetime.strptime(fecha_informe, '%d/%m/%Y')
    fecha_name = fecha.strftime('%Y-%m-%d ')
    nombre = "Informe de personal "+ " - " + str(fecha_name) +" - " +str(numero_registro_informe) 
    continiuacion_hr = str(numero_registro_informe) + "/"+ str(linea_mando)


    fecha_informe = self.ui.fecha_insitop.text() 
    fecha = datetime.strptime(fecha_informe, '%d/%m/%Y')

       
    unidad_toe = self.ui.sigla_ut_2.text()
    unidad_toe = unidad_toe.strip()
    query =  "select * from TOE_BATALLON where sigla like '{}%'".format(unidad_toe)
    
    datos_respuesta = select_dinamico_bd(self, TOE_BATALLON=query)

    efectivos_toe = select_dinamico_bd_personal(self, PERSONAL="*")
    ofi_toe = 0
    sub_toe = 0
    slp_toe = 0
    sl18_toe = 0
    sl12_toe = 0
    nombre_unidad = ""
    total_porcentaje=0
    for e in datos_respuesta[0]:


        # ofi_toe = int(e[0]) 

        ofi_toe = int(numero_diferente_cero(e[12]))
        sub_toe = int(numero_diferente_cero(e[21]))
        slp_toe = int(numero_diferente_cero(e[24]))
        sl18_toe = int(numero_diferente_cero(e[27]))
        sl12_toe = int(numero_diferente_cero(e[28]))
        nombre_unidad  = e[1]

    lambda_func = lambda x: x if x != 0 else ""

        #llenado de parte de personal en el informe de INSITOP
    context = { 

        'fecha_elaboracion' :  fecha_name,
        'registro_unidad' :     numero_registro_informe,
        'division': divisiones,
        'brigada': brigada,
        'unidad': nombre_unidad,
        'continuación_doc': continiuacion_hr,               
    }


    doc.render(context)

    import docx 
    
    doc2 = docx.Document() 
    
    doc2.add_heading('GeeksForGeeks', 0) 
    


    #codigo fuente para calcular los resultados de los cuadros que se exponen a contunuacion 


    ofi = 0
    sub = 0
    slp = 0
    sl18 = 0
    sl12 = 0
    h = 0
    m = 0

    ofi_agre = 0
    sub_agre = 0
    slp_agre = 0
    sl18_agre = 0
    sl12_agre = 0
    h_agre = 0
    m_agre = 0


    ofi_segre = 0
    sub_segre = 0
    slp_segre = 0
    sl18_segre = 0
    sl12_segre = 0
    h_segre = 0
    m_segre = 0

    ofi_organica = 0
    sub_organica = 0
    slp_organica = 0
    sl18_organica = 0
    sl12_organica = 0
    h_organica = 0
    m_organica = 0

    ofi_toe = 0
    sub_toe = 0
    slp_toe = 0
    sl18_toe = 0
    sl12_toe = 0


    ofi_novedades = 0
    sub_novedades = 0
    slp_novedades = 0
    sl18_novedades = 0
    sl12_novedades = 0
    h_novedades = 0
    m_novedades = 0

    ofi_code =0
    sub_code =0
    slp_code =0
    sl18_code =0
    sl12_code =0
    h_code =0
    m_code =0

    ofi_code_entrenamineto =0
    sub_code_entrenamineto =0
    slp_code_entrenamineto =0
    sl18_code_entrenamineto =0
    sl12_code_entrenamineto =0
    h_code_entrenamineto =0
    m_code_entrenamineto =0

    ofi_code_descanso =0
    sub_code_descanso =0
    slp_code_descanso =0
    sl18_code_descanso =0
    sl12_code_descanso =0
    h_code_descanso =0
    m_code_descanso =0

    ofi_pdamd =0
    sub_pdamd =0
    slp_pdamd =0
    sl18_pdamd =0
    sl12_pdamd =0
    h_pdamd =0
    m_pdamd =0

    
    ofi_pdmat =0
    sub_pdmat =0
    slp_pdmat =0
    sl18_pdmat =0
    sl12_pdmat =0
    h_pdmat =0
    m_pdmat =0

                
    ofi_comando = 0
    sub_comando = 0
    slp_comando = 0
    sl18_comando = 0
    sl12_comando = 0
    h_comando = 0
    m_comando = 0

    ofi_plana = 0
    sub_plana = 0
    slp_plana = 0
    sl18_plana = 0
    sl12_plana = 0
    h_plana = 0
    m_plana = 0
    comando = ""
    ejecutivo = ""


    datos_respuesta = select_dinamico_bd_personal(self, PERSONAL="*")
    for i in datos_respuesta[0]:

                        if i[14] == "OFI":
                                ofi = ofi + 1
                        elif i[14] == "SUB":
                                sub = sub +1
                        elif i[14] == "SLP":
                                slp = slp + 1
                        elif i[14] == "SL18":
                                sl18 = sl18 +1
                        elif i[14] == "SL12":
                                sl12 = sl12 +1

                        if i[13] == "MASCULINO":
                                h = h +1
                        else:
                                m = m +1
                        if i[9] == "NOVEDAD":
                                if i[14] == "OFI":
                                        ofi_novedades = ofi_novedades + 1
                                elif i[14] == "SUB":
                                        sub_novedades = sub_novedades +1
                                elif i[14] == "SLP":
                                        slp_novedades = slp_novedades + 1
                                elif i[14] == "SL18":
                                        sl18_novedades = sl18_novedades +1
                                elif i[14] == "SL12":
                                        sl12_novedades = sl12_novedades +1

                                if i[13] == "MASCULINO":
                                        h_novedades = h_novedades +1
                                else:
                                        m_novedades = m_novedades +1
                        else:
                                if i[10] == "OPERACIONES":
                                    if i[14] == "OFI":
                                            ofi_code = ofi_code + 1
                                    elif i[14] == "SUB":
                                            sub_code = sub_code +1
                                    elif i[14] == "SLP":
                                            slp_code = slp_code + 1
                                    elif i[14] == "SL18":
                                            sl18_code = sl18_code +1
                                    elif i[14] == "SL12":
                                            sl12_code = sl12_code +1

                                    if i[13] == "MASCULINO":
                                            h_code = h_code +1
                                    else:
                                            m_code = m_code +1

                                elif i[10] == "DESCANSO":
                                    if i[14] == "OFI":
                                            ofi_code_descanso = ofi_code_descanso + 1
                                    elif i[14] == "SUB":
                                            sub_code_descanso = sub_code_descanso +1
                                    elif i[14] == "SLP":
                                            slp_code_descanso = slp_code_descanso + 1
                                    elif i[14] == "SL18":
                                            sl18_code_descanso = sl18_code_descanso +1
                                    elif i[14] == "SL12":
                                            sl12_code_descanso = sl12_code_descanso +1

                                    if i[13] == "MASCULINO":
                                            h_code_descanso = h_code_descanso +1
                                    else:
                                            m_code_descanso = m_code_descanso +1
                                      
                                elif i[10] == "ENTRENAMIENTO":
                                    if i[14] == "OFI":
                                            ofi_code_entrenamineto = ofi_code_entrenamineto + 1
                                    elif i[14] == "SUB":
                                            sub_code_entrenamineto = sub_code_entrenamineto +1
                                    elif i[14] == "SLP":
                                            slp_code_entrenamineto = slp_code_entrenamineto + 1
                                    elif i[14] == "SL18":
                                            sl18_code_entrenamineto = sl18_code_entrenamineto +1
                                    elif i[14] == "SL12":
                                            sl12_code_entrenamineto = sl12_code_entrenamineto +1

                                    if i[13] == "MASCULINO":
                                            h_code_entrenamineto = h_code_entrenamineto +1
                                    else:
                                            m_code_entrenamineto = m_code_entrenamineto +1
                                      


                        if i[8] == "ORGÁNICA":
                                if i[14] == "OFI":
                                        ofi_organica = ofi_organica + 1
                                elif i[14] == "SUB":
                                        sub_organica = sub_organica +1
                                elif i[14] == "SLP":
                                        slp_organica = slp_organica + 1
                                elif i[14] == "SL18":
                                        sl18_organica = sl18_organica +1
                                elif i[14] == "SL12":
                                        sl12_organica = sl12_organica +1

                                if i[13] == "MASCULINO":
                                        h_organica = h_organica +1
                                else:
                                        m_organica = m_organica +1

                        elif i[8] == "SEGREGADO":
                                if i[14] == "OFI":
                                        ofi_segre = ofi_segre + 1
                                elif i[14] == "SUB":
                                        sub_segre = sub_segre +1
                                elif i[14] == "SLP":
                                        slp_segre = slp_segre + 1
                                elif i[14] == "SL18":
                                        sl18_segre = sl18_segre +1
                                elif i[14] == "SL12":
                                        sl12_segre = sl12_segre +1

                                if i[13] == "MASCULINO":
                                        h_segre = h_segre +1
                                else:
                                        m_segre = m_segre +1

                        else :
                                if i[14] == "OFI":
                                        ofi_agre = ofi_agre + 1
                                elif i[14] == "SUB":
                                        sub_agre = sub_agre +1
                                elif i[14] == "SLP":
                                        slp_agre = slp_agre + 1
                                elif i[14] == "SL18":
                                        sl18_agre = sl18_agre +1
                                elif i[14] == "SL12":
                                        sl12_agre = sl12_agre +1

                                if i[13] == "MASCULINO":
                                        h_agre = h_agre +1
                                else:
                                        m_agre = m_agre +1

                        if i[11] == "PMAD":
                                if i[14] == "OFI":
                                        ofi_pdamd = ofi_pdamd + 1
                                elif i[14] == "SUB":
                                        sub_pdamd = sub_pdamd +1
                                elif i[14] == "SLP":
                                        slp_pdamd = slp_pdamd + 1
                                elif i[14] == "SL18":
                                        sl18_pdamd = sl18_pdamd +1
                                elif i[14] == "SL12":
                                        sl12_pdamd = sl12_pdamd +1

                                if i[13] == "MASCULINO":
                                        h_pdamd = h_pdamd +1
                                else:
                                        m_pdamd = m_pdamd +1
                        elif i[11] == "PDMAT":
                                if i[14] == "OFI":
                                        ofi_pdmat = ofi_pdmat + 1
                                elif i[14] == "SUB":
                                        sub_pdmat = sub_pdmat +1
                                elif i[14] == "SLP":
                                        slp_pdmat = slp_pdmat + 1
                                elif i[14] == "SL18":
                                        sl18_pdmat = sl18_pdmat +1
                                elif i[14] == "SL12":
                                        sl12_pdmat = sl12_pdmat +1

                                if i[13] == "MASCULINO":
                                        h_pdmat = h_pdmat +1
                                else:
                                        m_pdmat = m_pdmat +1

                                            
                        if i[6] == "CDO" and i[7]  == "6":
                            if i[14] == "OFI":
                                    ofi_comando = ofi_comando + 1
                            elif i[14] == "SUB":
                                    sub_comando = sub_comando +1
                            elif i[14] == "SLP":
                                    slp_comando = slp_comando + 1
                            elif i[14] == "SL18":
                                    sl18_comando = sl18_comando +1
                            elif i[14] == "SL12":
                                    sl12_comando = sl12_comando +1

                            if i[13] == "MASCULINO":
                                    h_comando = h_comando +1
                            else:
                                    m_comando = m_comando +1

                        elif i[6] == "CDO" and i[7]  == "5" or i[6] == "CDO" and i[7]  == "3"  or i[6] == "PLM":
                            if i[14] == "OFI":
                                    ofi_plana = ofi_plana + 1
                            elif i[14] == "SUB":
                                    sub_plana = sub_plana +1
                            elif i[14] == "SLP":
                                    slp_plana = slp_plana + 1
                            elif i[14] == "SL18":
                                    sl18_plana = sl18_plana +1
                            elif i[14] == "SL12":
                                    sl12_plana = sl12_plana +1

                            if i[13] == "MASCULINO":
                                    h_plana = h_plana +1
                            else:
                                    m_plana = m_plana + 1




    table3 = doc.add_table(rows=1, cols=15) 
    table3.style = 'firmas' 
    row = table3.rows[0].cells 

    row[14].merge(row[0])
    row[0].paragraphs[0].add_run(str("INSITROP")).bold = True




    table = doc.add_table(rows=1, cols=9) 
    table.style = 'unidades' 
    row = table.rows[0].cells 

    p =row[1].merge(row[0])
    row[0].paragraphs[0].add_run("PERSONAL").bold = True
    p = row[0].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    
    row[2].paragraphs[0].add_run("OFI").bold = True
    p = row[2].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[3].paragraphs[0].add_run("SUB").bold = True
    p = row[3].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[4].paragraphs[0].add_run("SLP").bold = True
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[5].paragraphs[0].add_run("SL18").bold = True
    p = row[5].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[6].paragraphs[0].add_run("SL12").bold = True
    p = row[6].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[7].paragraphs[0].add_run("H").bold = True
    p = row[7].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[8].paragraphs[0].add_run("M").bold = True
    p = row[8].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    # 
    
    row = table.add_row().cells 
    row[1].merge(row[0])
    row[0].text = str("TOE")
    row[2].text = str(ofi_toe)
    p = row[2].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[3].text = str(sub_toe)
    p = row[3].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[4].text = str(slp_toe)
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[5].text = str(sl18_toe)
    p = row[5].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[6].text = str(sl12_toe)
    p = row[6].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[7].text = str("")
    p = row[7].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[8].text = str("")
    p = row[8].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

    row = table.add_row().cells 
    row[1].merge(row[0])
    row[0].text = str("EFECTIVOS")
    row[2].text = str(ofi)
    p = row[2].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[3].text = str(sub)
    p = row[3].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[4].text = str(slp)
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[5].text = str(sl18)
    p = row[5].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[6].text = str(sl12)
    p = row[6].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[7].text = str(h)
    p = row[7].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[8].text = str(m)
    p = row[8].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

    
    row = table.add_row().cells 
    row[1].merge(row[0])
    row[0].text = str("SEGREGADOS")
    row[2].text = str(ofi_segre)
    p = row[2].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[3].text = str(sub_segre)
    p = row[3].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[4].text = str(slp_segre)
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[5].text = str(sl18_segre)
    p = row[5].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[6].text = str(sl12_segre)
    p = row[6].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[7].text = str(h_segre)
    p = row[7].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[8].text = str(m_segre)
    p = row[8].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
        
    row = table.add_row().cells 
    row[1].merge(row[0])
    row[0].text = str("AGREGACIONES")
    row[2].text = str(ofi_agre)
    p = row[2].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[3].text = str(sub_agre)
    p = row[3].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[4].text = str(slp_agre)
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[5].text = str(sl18_agre)
    p = row[5].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[6].text = str(sl12_agre)
    p = row[6].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[7].text = str(h_agre)
    p = row[7].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[8].text = str(m_agre)
    p = row[8].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

            
    row = table.add_row().cells 
    row[1].merge(row[0])
    row[0].text = str("NOVEDADES")
    row[2].text = str(ofi_novedades)
    p = row[2].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[3].text = str(sub_novedades)
    p = row[3].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[4].text = str(slp_novedades)
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[5].text = str(sl18_novedades)
    p = row[5].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[6].text = str(sl12_novedades)
    p = row[6].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[7].text = str(h_novedades)
    p = row[7].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[8].text = str(m_novedades)
    p = row[8].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
                
    row = table.add_row().cells 
    row[1].merge(row[0])
    row[0].text = str("AREA OPERACIONES")
    row[2].text = str(ofi_code)
    p = row[2].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[3].text = str(sub_code)
    p = row[3].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[4].text = str(slp_code)
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[5].text = str(sl18_code)
    p = row[5].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[6].text = str(sl12_code)
    p = row[6].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[7].text = str(h_code)
    p = row[7].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[8].text = str(m_code)
    p = row[8].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
                    
    row = table.add_row().cells 
    row[1].merge(row[0])
    row[0].text = str("ENTRENAMIENTO")
    row[2].text = str(ofi_code_entrenamineto)
    p = row[2].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[3].text = str(sub_code_entrenamineto)
    p = row[3].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[4].text = str(slp_code_entrenamineto)
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[5].text = str(sl18_code_entrenamineto)
    p = row[5].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[6].text = str(sl12_code_entrenamineto)
    p = row[6].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[7].text = str(h_code_entrenamineto)
    p = row[7].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[8].text = str(m_code_entrenamineto)
    p = row[8].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

                        
    row = table.add_row().cells 
    row[1].merge(row[0])
    row[0].text = str("DESCANSO")
    row[2].text = str(ofi_code_descanso)
    p = row[2].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[3].text = str(sub_code_descanso)
    p = row[3].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[4].text = str(slp_code_descanso)
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[5].text = str(sl18_code_descanso)
    p = row[5].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[6].text = str(sl12_code_descanso)
    p = row[6].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[7].text = str(h_code_descanso)
    p = row[7].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[8].text = str(m_code_descanso)
    p = row[8].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
                            
    row = table.add_row().cells 
    row[1].merge(row[0])
    row[0].text = str("PDMAD")
    row[2].text = str(ofi_pdamd)
    p = row[2].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[3].text = str(sub_pdamd)
    p = row[3].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[4].text = str(slp_pdamd)
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[5].text = str(sl18_pdamd)
    p = row[5].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[6].text = str(sl12_pdamd)
    p = row[6].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[7].text = str(h_pdamd)
    p = row[7].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[8].text = str(m_pdamd)
    p = row[8].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

                            
    row = table.add_row().cells 
    row[1].merge(row[0])
    row[0].text = str("PDMAT")
    row[2].text = str(ofi_pdmat)
    p = row[2].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[3].text = str(sub_pdmat)
    p = row[3].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[4].text = str(slp_pdmat)
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[5].text = str(sl18_pdmat)
    p = row[5].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[6].text = str(sl12_pdmat)
    p = row[6].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[7].text = str(h_pdmat)
    p = row[7].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[8].text = str(m_pdmat)
    p = row[8].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

                            
    row = table.add_row().cells 
    row[1].merge(row[0])
    row[0].text = str("CDO")
    row[2].text = str(ofi_comando)
    p = row[2].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[3].text = str(sub_comando)
    p = row[3].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[4].text = str(slp_comando)
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[5].text = str(sl18_comando)
    p = row[5].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[6].text = str(sl12_comando)
    p = row[6].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[7].text = str(h_comando)
    p = row[7].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[8].text = str(m_comando)
    p = row[8].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

                            
    row = table.add_row().cells 
    row[1].merge(row[0])
    row[0].text = str("PLANA MAYOR")
    row[2].text = str(ofi_plana)
    p = row[2].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[3].text = str(sub_plana)
    p = row[3].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[4].text = str(slp_plana)
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[5].text = str(sl18_plana)
    p = row[5].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[6].text = str(sl12_plana)
    p = row[6].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[7].text = str(h_plana)
    p = row[7].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
    row[8].text = str(m_plana)
    p = row[8].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

    parrafo = doc.add_paragraph()
               
    table3 = doc.add_table(rows=1, cols=15) 
    table3.style = 'firmas' 
    row = table3.rows[0].cells 

    row[14].merge(row[0])
    row[0].paragraphs[0].add_run(str("RELACIÓN DE PELOTONES")).bold = True  

    table = doc.add_table(rows=1, cols=30) 
    table.style = 'unidades' 
    row = table.rows[0].cells 




    #cantidad de pelotones 
        
    informacion =" SELECT distinct DIVISION, BRIGADA, BATALLON, COMPANIA, PELOTON, GRD,  APELLIDOS_NOMBRES, ACTIVIDAD, RELACION_MANDO, SEXO, ESCALAFON  FROM PERSONAL where SITUACION like 'CODE' and COMPANIA not like 'PLM'  AND CARGO = 'CDTE PELOTON' ORDER BY COMPANIA ASC"

    datos = b_d_i.INFORMACION_PEL(self, informacion)

    fila = 0
    pel_operaciones = 0

    pel_entrenamiento = 0

    pel_descanso = 0

    cdt_ofi = 0
    cdt_sub = 0
    cdt_h = 0
    cdt_m = 0

    pel_organico = 0
    pelotones_agr = 0

    for registro in datos:
            if registro[8] == "ORGÁNICA":
                    pel_organico =  pel_organico + 1

            elif registro[8] != "ORGÁNICA" and registro[8] != "SEGREGADO":
                    pelotones_agr = pelotones_agr +1
                    
            if registro[7] == "OPERACIONES":
                 pel_operaciones = pel_operaciones +1
            elif registro[7] == "DESCANSO":
                 pel_descanso = pel_descanso +1
            elif registro[7] == "ENTRENAMIENTO":
                 pel_entrenamiento = pel_entrenamiento +1

            if registro[10] == "OFI":
                 cdt_ofi = cdt_ofi +1
            else:
                 cdt_sub = cdt_sub +1

            if registro[9] == "MASCULINO":
                 cdt_h = cdt_h +1
            else:
                 cdt_m = cdt_m +1




#pelotones linea 1

    row[3].merge(row[0])
    row[0].paragraphs[0].add_run("CDTE OFI").bold = True
    p = row[0].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[5].merge(row[4])
    row[4].text = str(cdt_ofi)
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER    
    
    row[9].merge(row[6])
    row[6].paragraphs[0].add_run("CDTE SUB").bold = True
    p = row[6].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  
   
        
    row[11].merge(row[10])
    row[10].text = str(cdt_sub)
    p = row[10].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER     
            
    row[15].merge(row[12])
    row[12].paragraphs[0].add_run("CDTE H").bold = True
    p = row[12].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[17].merge(row[16])
    row[16].text = str(cdt_h)
    p = row[16].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[21].merge(row[18])
    row[18].paragraphs[0].add_run("CDTE M").bold = True 
    p = row[18].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  
    
    row[23].merge(row[22])
    row[22].text = str(cdt_m)
    p = row[22].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[27].merge(row[24])
    row[24].paragraphs[0].add_run("TOTAL PELOTONES").bold = True 
    p = row[24].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  
    total = pel_operaciones + pel_entrenamiento + pel_descanso
    row[29].merge(row[28])
    row[28].text = str(total)
    p = row[28].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    #pelotones linea2
    row = table.add_row().cells 

    row[3].merge(row[0])
    row[0].paragraphs[0].add_run("PELOTONES AREÁ").bold = True
    p = row[0].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[5].merge(row[4])
    row[4].text = str(pel_operaciones)
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER    
    
    row[9].merge(row[6])
    row[6].paragraphs[0].add_run("ENTRENAMIENTOS").bold = True
    p = row[6].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  
   
        
    row[11].merge(row[10])
    row[10].text = str(pel_entrenamiento)
    p = row[10].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER     
            
    row[15].merge(row[12])
    row[12].paragraphs[0].add_run("DESCANSOS").bold = True
    p = row[12].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[17].merge(row[16])
    row[16].text = str(pel_descanso)
    p = row[16].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[21].merge(row[18])
    row[18].paragraphs[0].add_run("PEL ORGANICOS").bold = True 
    p = row[18].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  
    
    row[23].merge(row[22])
    row[22].text = str(pel_organico)
    p = row[22].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[27].merge(row[24])
    row[24].paragraphs[0].add_run("PEL AGREGADOS").bold = True 
    p = row[24].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[29].merge(row[28])
    row[28].text = str(pelotones_agr)
    p = row[28].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  



#calcular la cantidad de exdes
    informacion =" SELECT distinct DIVISION, BRIGADA, BATALLON, COMPANIA, PELOTON, GRD,  APELLIDOS_NOMBRES, ACTIVIDAD, RELACION_MANDO, SEXO, ESCALAFON  FROM PERSONAL where SITUACION like 'CODE' and COMPANIA not like 'PLM'  AND CARGO = 'CDTE EXDE' ORDER BY COMPANIA ASC"
    datos = b_d_i.INFORMACION_PEL(self, informacion)

    exde_organico = 0
    exde_agregados = 0
    for registro in datos:
            if registro[8] == "ORGÁNICA":
                    exde_organico =  exde_organico + 1
                    # print(str(pel_organico)+str(". ")+x[2]+str(" - ")+x[3]+str(" - ")+x[4] +str(": ")+x[5])

            elif registro[8] != "ORGÁNICA" and registro[8] != "SEGREGADO":
                    exde_agregados = exde_agregados +1

    parrafo = doc.add_paragraph()
               
    table3 = doc.add_table(rows=1, cols=15) 
    table3.style = 'firmas' 
    row = table3.rows[0].cells 

    row[14].merge(row[0])
    row[0].paragraphs[0].add_run(str("EQUIDPOS EXDES")).bold = True  

    table = doc.add_table(rows=1, cols=6) 
    table.style = 'unidades' 
    row = table.rows[0].cells 


    row[0].merge(row[0])
    row[0].paragraphs[0].add_run("EXDE ORGANICOS").bold = True
    p = row[0].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[1].merge(row[1])
    row[1].text = str(exde_organico)
    p = row[1].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER    
    
    row[2].merge(row[2])
    row[2].paragraphs[0].add_run("EXDE SEGREGADOS").bold = True
    p = row[2].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  
   
        
    row[3].merge(row[3])
    row[3].text = str(exde_agregados)
    p = row[3].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER     
    
    row[4].merge(row[4])
    row[4].paragraphs[0].add_run("TOTAL EXDES").bold = True
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  
    total = exde_organico + exde_agregados
    row[5].merge(row[5])
    row[5].text = str(total)
    p = row[5].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

#calcular la cantidad de exdes
    informacion =" SELECT distinct operacion, lineas_esfuerzo, no_ordop, ordop  FROM EFECTIVOS_PEL "
    datos = b_d_i.INFORMACION_PEL(self, informacion)

    OFENSIVAS = 0
    DEFENSIVAS = 0
    ESTABILIDAD = 0
    COMAC = 0

    LOE_TENERIFE = 0
    LOE_JUNNIN = 0
    LOE_BARBUKA = 0
    LOE_GAMEZA = 0
    LOE_CIENEGA = 0
    LOE_AMAZONIA = 0
    LOO_BOMBONA = 0
    LOO_MUTIS = 0
    LOO_PAYA = 0
    LOO_BOYACA = 0
    LOO_PANTANO_DE_VARGAS = 0
    LOO_TARAPACA = 0
                    
                    
                    
                    
    for registro in datos:

            if registro[0] == "OPERACIÓN OFENSIVAS":
                    OFENSIVAS =  OFENSIVAS + 1
            elif registro[0] == "OPERACIÓN DEFENSIVAS":
                    DEFENSIVAS =  DEFENSIVAS + 1
            elif registro[0] == "OPERACIÓN ESTABILIDAD":
                    ESTABILIDAD =  ESTABILIDAD + 1
            elif registro[0] == "OPERACIÓN COMAC":
                    COMAC =  COMAC + 1
                    # print(str(pel_organico)+str(". ")+x[2]+str(" - ")+x[3]+str(" - ")+x[4] +str(": ")+x[5])

            if registro[1] == "LOE TENERIFE":
                    LOE_TENERIFE = LOE_TENERIFE +1
            elif registro[1] == "LOE JUNÍN":
                    LOE_JUNNIN = LOE_JUNNIN +1
            elif registro[1] == "LOE BÁRBULA":
                    LOE_BARBUKA = LOE_BARBUKA +1
            elif registro[1] == "LOE GÁMEZA":
                    LOE_GAMEZA = LOE_GAMEZA +1
            elif registro[1] == "LOE CIÉNAGA":
                    LOE_CIENEGA = LOE_CIENEGA +1
            elif registro[1] == "LOE AMAZONÍA":
                    LOE_AMAZONIA = LOE_AMAZONIA +1
            elif registro[1] == "LOO BOMBONÁ":
                    LOO_BOMBONA = LOO_BOMBONA +1
            elif registro[1] == "LOO MUTIS":
                    LOO_MUTIS = LOO_MUTIS +1
            elif registro[1] == "LOO PAYA":
                    LOO_PAYA = LOO_PAYA +1
            elif registro[1] == "LOO BOYACÁ":
                    LOO_BOYACA = LOO_BOYACA +1
            elif registro[1] == "LOO PANTANO DE VARGAS":
                    LOO_PANTANO_DE_VARGAS = LOO_PANTANO_DE_VARGAS +1
            elif registro[1] == "LOO TARAPACÁ":
                    LOO_TARAPACA = LOO_TARAPACA +1

    parrafo = doc.add_paragraph()
               
    table3 = doc.add_table(rows=1, cols=15) 
    table3.style = 'firmas' 
    row = table3.rows[0].cells 

    row[14].merge(row[0])
    row[0].paragraphs[0].add_run(str("ORDENES DE OPERACIONES")).bold = True  

    table = doc.add_table(rows=1, cols=30) 
    table.style = 'unidades' 
    row = table.rows[0].cells 


    #pelotones linea 1

    row[3].merge(row[0])
    row[0].paragraphs[0].add_run("OFENSIVAS").bold = True
    p = row[0].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[5].merge(row[4])
    row[4].text = str(OFENSIVAS)
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER    
    
    row[9].merge(row[6])
    row[6].paragraphs[0].add_run("DEFENSIVAS").bold = True
    p = row[6].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  
   
        
    row[11].merge(row[10])
    row[10].text = str(DEFENSIVAS)
    p = row[10].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER     
            
    row[15].merge(row[12])
    row[12].paragraphs[0].add_run("ESTABILIDAD").bold = True
    p = row[12].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[17].merge(row[16])
    row[16].text = str(ESTABILIDAD)
    p = row[16].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[21].merge(row[18])
    row[18].paragraphs[0].add_run("COMAC").bold = True 
    p = row[18].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  
    
    row[23].merge(row[22])
    row[22].text = str(COMAC)
    p = row[22].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    total =OFENSIVAS + DEFENSIVAS + ESTABILIDAD + COMAC
    row[27].merge(row[24])
    row[24].paragraphs[0].add_run("TOTAL ORDOP").bold = True 
    p = row[24].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[29].merge(row[28])
    row[28].text = str(total)
    p = row[28].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER 

    parrafo = doc.add_paragraph()

    table3 = doc.add_table(rows=1, cols=15) 
    table3.style = 'firmas' 
    row = table3.rows[0].cells 

    row[14].merge(row[0])
    row[0].paragraphs[0].add_run(str("LINEAS DE ESFUERZO")).bold = True  

    table = doc.add_table(rows=1, cols=30) 
    table.style = 'unidades' 
    row = table.rows[0].cells 

    row[3].merge(row[0])
    row[0].paragraphs[0].add_run("LOE TENERIFE").bold = True
    p = row[0].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[5].merge(row[4])
    row[4].text = str(LOE_TENERIFE)
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER    
    
    row[9].merge(row[6])
    row[6].paragraphs[0].add_run("LOE JUNÍN").bold = True
    p = row[6].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  
   
        
    row[11].merge(row[10])
    row[10].text = str(LOE_JUNNIN)
    p = row[10].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER     
            
    row[15].merge(row[12])
    row[12].paragraphs[0].add_run("LOE BÁRBULA").bold = True
    p = row[12].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[17].merge(row[16])
    row[16].text = str(LOE_BARBUKA)
    p = row[16].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[21].merge(row[18])
    row[18].paragraphs[0].add_run("LOE GÁMEZA").bold = True 
    p = row[18].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  
    
    row[23].merge(row[22])
    row[22].text = str(LOE_GAMEZA)
    p = row[22].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[27].merge(row[24])
    row[24].paragraphs[0].add_run("LOE CIÉNAGA").bold = True 
    p = row[24].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[29].merge(row[28])
    row[28].text = str(LOE_CIENEGA)
    p = row[28].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  


    row = table.add_row().cells 

    row[3].merge(row[0])
    row[0].paragraphs[0].add_run("LOE AMAZONÍA").bold = True
    p = row[0].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[5].merge(row[4])
    row[4].text = str(LOE_AMAZONIA)
    p = row[4].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER    
    
    row[9].merge(row[6])
    row[6].paragraphs[0].add_run("LOO BOMBONÁ").bold = True
    p = row[6].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  
   
        
    row[11].merge(row[10])
    row[10].text = str(LOO_BOMBONA)
    p = row[10].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER     
            
    row[15].merge(row[12])
    row[12].paragraphs[0].add_run("LOO MUTIS").bold = True
    p = row[12].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[17].merge(row[16])
    row[16].text = str(LOO_MUTIS)
    p = row[16].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[21].merge(row[18])
    row[18].paragraphs[0].add_run("LOO PAYA").bold = True 
    p = row[18].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  
    
    row[23].merge(row[22])
    row[22].text = str(LOO_PAYA)
    p = row[22].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[27].merge(row[24])
    row[24].paragraphs[0].add_run("LOO BOYACÁ").bold = True 
    p = row[24].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[29].merge(row[28])
    row[28].text = str(LOO_BOYACA)
    p = row[28].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER 


    row = table.add_row().cells 

    row[6].merge(row[0])
    row[0].paragraphs[0].add_run("LOO PANTANO DE VARGAS").bold = True
    p = row[0].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

    row[8].merge(row[7])
    row[7].text = str(LOO_PANTANO_DE_VARGAS)
    p = row[7].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER    
    
    row[13].merge(row[9])
    row[9].paragraphs[0].add_run("LOO TARAPACÁ").bold = True
    p = row[9].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  
   
    row[15].merge(row[14])
    row[14].text = str(LOO_TARAPACA)
    p = row[14].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER    
 
    row[20].merge(row[16])
    row[16].paragraphs[0].add_run("TOTAL LOE").bold = True
    p = row[16].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  
   

    TOTAL_LOE = LOE_TENERIFE + LOE_JUNNIN + LOE_BARBUKA + LOE_GAMEZA + LOE_CIENEGA + LOE_AMAZONIA
    TOTAL_LOO = LOO_BOMBONA + LOO_MUTIS + LOO_PAYA + LOO_BOYACA + LOO_PANTANO_DE_VARGAS + LOO_TARAPACA
    row[22].merge(row[21])
    row[21].text = str(TOTAL_LOE)
    p = row[21].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER   



    row[27].merge(row[23])
    row[23].paragraphs[0].add_run("TOTAL LOO").bold = True
    p = row[23].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  
           
    row[29].merge(row[28])
    row[28].text = str(TOTAL_LOO)
    p = row[28].paragraphs[0]
    p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  


    parrafo = doc.add_paragraph()
    parrafo = doc.add_paragraph()



    fecha_informe = self.ui.fecha_insitop.text()
    date_fecha_inicial=datetime.strptime(fecha_informe, '%d/%m/%Y') 

    query = """SELECT * FROM INSITOP WHERE FECHA >= '{}' and FECHA <= '{}' ORDER BY CP ASC """.format(date_fecha_inicial,date_fecha_inicial)

    unidades_p = b_d_i.INFORMACION_PEL(self, query)
    
    print(unidades_p)
    
    if unidades_p !=[]:
        
        table3 = doc.add_table(rows=1, cols=15) 
        table3.style = 'firmas' 
        row = table3.rows[0].cells 

        row[14].merge(row[0])
        row[0].paragraphs[0].add_run(str("INFORME DE MOVIMIENTO")).bold = True  

        table = doc.add_table(rows=1, cols=30) 
        table.style = 'resultados' 
        row = table.rows[0].cells
  
        for unidad in unidades_p:

                nombre_unidad_seguimiento = str(unidad[7])+" - "+str(unidad[8])+str(unidad[9])
                nombre_cdte = str(unidad[12])+". "+str(unidad[13])+" ("+str(unidad[14])+") - "+str(unidad[15])
                dept = str(unidad[24])
                mpio = str(unidad[25])
                lugar = str(unidad[26])
                distancia = str(unidad[83]) +" m"
                coordenadas = str(unidad[16])+" "+str(unidad[17])+"° "+str(unidad[18])+"' "+str(unidad[19])+"'' - "+str(unidad[20])+" "+str(unidad[21])+"° "+str(unidad[22])+"' "+str(unidad[23])+"''"
                
                area ="AREÁ OP. "+ str(unidad[27])+"-"+str(unidad[28])+"-"+str(unidad[29])+"-"+str(unidad[30])+"-"+str(unidad[31])+" ("+str(unidad[32])+"-"+str(unidad[33])+")"

                efectivos ="EFECTIVOS "+ str(unidad[34])+"-"+str(unidad[35])+"-"+str(unidad[36])+"-"+str(unidad[37])+"-"+str(unidad[38])+" ("+str(unidad[39])+"-"+str(unidad[40])+")"
                
                disponibles = "DISP. "+ str(unidad[41])+"-"+str(unidad[42])+"-"+str(unidad[43])+"-"+str(unidad[44])+"-"+str(unidad[45])+" ("+str(unidad[46])+"-"+str(unidad[47])+")"
                        
                noveades = "NOV "+ str(unidad[48])+"-"+str(unidad[49])+"-"+str(unidad[50])+"-"+str(unidad[51])+"-"+str(unidad[52])+" ("+str(unidad[53])+"-"+str(unidad[54])+")"

                exde = "EXDE "+ str(unidad[55])+"-"+str(unidad[56])+"-"+str(unidad[57])+"-"+str(unidad[58])+"-"+str(unidad[59])+" ("+str(unidad[60])+"-"+str(unidad[61])+") TIENE "+str(unidad[62])

                exde_CDTE = "cdte exde "+str(unidad[63])+". "+str(unidad[64])+" ("+str(unidad[65])+")"

                ordop = str(unidad[74])+" "+str(unidad[75])

                operacion = str(unidad[69])
                tarea = str(unidad[70])

                tecnica = str(unidad[71])
                loo_loe = str(unidad[72])
                relacion_mando = str(unidad[10])
                relacion_mando_doc = "DOC. AGRE "+str(unidad[11])

                row[2].merge(row[0])
                row[0].paragraphs[0].add_run(nombre_unidad_seguimiento).bold = True
                p = row[0].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  
                
                row[11].merge(row[3])
                row[3].paragraphs[0].add_run(nombre_cdte).bold = True
                p = row[3].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
                        
                row[16].merge(row[12])
                row[12].paragraphs[0].add_run(coordenadas).bold = True
                p = row[12].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

                row[22].merge(row[17])
                row[17].paragraphs[0].add_run(dept).bold = True
                p = row[17].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
                
                
                row[29].merge(row[23])
                row[23].paragraphs[0].add_run(mpio).bold = True
                p = row[23].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER    
                
                
                row = table.add_row().cells
                row[8].merge(row[0])
                row[0].paragraphs[0].add_run(lugar).bold = True
                p = row[0].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  
                
                row[11].merge(row[9])
                row[9].paragraphs[0].add_run(distancia).bold = True
                p = row[9].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

                row[16].merge(row[12])
                row[12].paragraphs[0].add_run(area).bold = True
                p = row[12].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER 

                row[21].merge(row[17])
                row[17].paragraphs[0].add_run(efectivos).bold = True
                p = row[17].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER 
                
        
                row[25].merge(row[22])
                row[22].paragraphs[0].add_run(disponibles).bold = True
                p = row[22].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER 

                row[29].merge(row[26])
                row[26].paragraphs[0].add_run(noveades).bold = True
                p = row[26].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER 
                
                row = table.add_row().cells
                
                row[9].merge(row[0])
                row[0].paragraphs[0].add_run(exde_CDTE).bold = True
                p = row[0].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

                row[15].merge(row[10])
                row[10].paragraphs[0].add_run(exde).bold = True
                p = row[10].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER 
                
                row[19].merge(row[16])
                row[16].paragraphs[0].add_run(ordop).bold = True
                p = row[16].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER 

                row[24].merge(row[20])
                row[20].paragraphs[0].add_run(operacion).bold = True
                p = row[20].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER 

                row[29].merge(row[25])
                row[25].paragraphs[0].add_run(tarea).bold = True
                p = row[25].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER 

                row = table.add_row().cells
                
                row[13].merge(row[0])
                row[0].paragraphs[0].add_run(tecnica).bold = True
                p = row[0].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER 

                row[19].merge(row[14])
                row[14].paragraphs[0].add_run(loo_loe).bold = True
                p = row[14].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER 

                row[23].merge(row[20])
                row[20].paragraphs[0].add_run(relacion_mando).bold = True
                p = row[20].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER 

                row[29].merge(row[24])
                row[24].paragraphs[0].add_run(relacion_mando_doc).bold = True
                p = row[24].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER 

                row = table.add_row().cells

                code_fecha = "UNIDAD EN "+str(unidad[73])+str(" DESDE ")+str(unidad[76])
                dia_code = str(unidad[79])
                OBSERVACION = str(unidad[82])
                
                row[9].merge(row[0])
                row[0].paragraphs[0].add_run(code_fecha).bold = True
                p = row[0].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER 
        
                row[15].merge(row[10])
                row[10].paragraphs[0].add_run(dia_code).bold = True
                p = row[10].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER 
        
                row[29].merge(row[16])
                row[16].paragraphs[0].add_run(OBSERVACION).bold = True
                p = row[16].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER 
        
                parrafo = doc.add_paragraph()

        

#calcular la cantidad de exdes

    informacion =" SELECT distinct DIVISION, BRIGADA, BATALLON, COMPANIA, PELOTON, GRD,  APELLIDOS_NOMBRES, ACTIVIDAD, UBICACION, RELACION_MANDO, SEXO, TELEFONO, ESCALAFON  FROM PERSONAL where SITUACION like 'CODE' and COMPANIA not like 'PLM'  AND CARGO = 'CDTE PELOTON' and ACTIVIDAD = 'ENTRENAMIENTO' ORDER BY COMPANIA ASC"
    datos = b_d_i.INFORMACION_PEL(self, informacion)
    
    if datos !=[]:

        parrafo = doc.add_paragraph()
                
        table3 = doc.add_table(rows=1, cols=15) 
        table3.style = 'firmas' 
        row = table3.rows[0].cells 

        row[14].merge(row[0])
        row[0].paragraphs[0].add_run(str("PELOTONES EN REENTRENAMIENTO")).bold = True  

        table = doc.add_table(rows=1, cols=20) 
        table.style = 'resultados' 
        row = table.rows[0].cells
        # DESCANSO
        NUMERO = len(datos)

        for unidad in datos:
                nombre_unidad_seguimiento = str(unidad[2])+" - "+str(unidad[3])+str(unidad[4])
                nombre_cdte = str(unidad[5])+". "+str(unidad[6])+" ("+str(unidad[11])+")"
                ubicacion = str(unidad[8])
                relacion = str(unidad[9])

                row[2].merge(row[0])
                row[0].paragraphs[0].add_run(nombre_unidad_seguimiento).bold = True
                p = row[0].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

                row[11].merge(row[3])
                row[3].paragraphs[0].add_run(nombre_cdte).bold = True
                p = row[3].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

                row[15].merge(row[12])
                row[12].paragraphs[0].add_run(ubicacion).bold = True
                p = row[12].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

                row[19].merge(row[16])
                row[16].paragraphs[0].add_run(relacion).bold = True
                p = row[16].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER
                NUMERO = NUMERO -1
                if NUMERO > 0:    
                        row = table.add_row().cells
                
     
                       
                        
        
    informacion =" SELECT distinct DIVISION, BRIGADA, BATALLON, COMPANIA, PELOTON, GRD,  APELLIDOS_NOMBRES, ACTIVIDAD, UBICACION, RELACION_MANDO, SEXO, TELEFONO, ESCALAFON  FROM PERSONAL where SITUACION like 'CODE' and COMPANIA not like 'PLM'  AND CARGO = 'CDTE PELOTON' and ACTIVIDAD = 'DESCANSO' ORDER BY COMPANIA ASC"
    datos = b_d_i.INFORMACION_PEL(self, informacion)
    
    if datos !=[]:

        parrafo = doc.add_paragraph()
                
        table3 = doc.add_table(rows=1, cols=15) 
        table3.style = 'firmas' 
        row = table3.rows[0].cells 

        row[14].merge(row[0])
        row[0].paragraphs[0].add_run(str("PELOTONES EN DESCANSO")).bold = True  

        table = doc.add_table(rows=1, cols=20) 
        table.style = 'resultados' 
        row = table.rows[0].cells
        NUMERO = len(datos)

        for unidad in datos:
                nombre_unidad_seguimiento = str(unidad[2])+" - "+str(unidad[3])+str(unidad[4])
                nombre_cdte = str(unidad[5])+". "+str(unidad[6])+" ("+str(unidad[11])+")"
                ubicacion = str(unidad[8])
                relacion = str(unidad[9])

                row[2].merge(row[0])
                row[0].paragraphs[0].add_run(nombre_unidad_seguimiento).bold = True
                p = row[0].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER  

                row[11].merge(row[3])
                row[3].paragraphs[0].add_run(nombre_cdte).bold = True
                p = row[3].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

                row[15].merge(row[12])
                row[12].paragraphs[0].add_run(ubicacion).bold = True
                p = row[12].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

                row[19].merge(row[16])
                row[16].paragraphs[0].add_run(relacion).bold = True
                p = row[16].paragraphs[0]
                p.alignment = docx.enum.text.WD_PARAGRAPH_ALIGNMENT.CENTER

                NUMERO = NUMERO -1
                if NUMERO > 0:    
                        row = table.add_row().cells
                        
                        
    try:
        
        parrafo = doc.add_paragraph()
        parrafo = doc.add_paragraph()
        parrafo = doc.add_paragraph()
        parrafo = doc.add_paragraph()
        parrafo = doc.add_paragraph()

        query = "SELECT * FROM PERSONAL WHERE CARGO like 'CDTE BATALLON'"
        efectivos_toe = b_d_i.INFORMACION_PEL(self, query)
    # datos_respuesta = datos_base_datos(self,comando="*")

        for i in efectivos_toe:
                #comando
                #comando
                cdte = str(i[0])+". "+str(i[1])

        query = "SELECT * FROM PERSONAL WHERE CARGO like 'OFICIAL DE OPERACIONES'"
        efectivos_toe_OFI = b_d_i.INFORMACION_PEL(self, query)
        for i in efectivos_toe_OFI:


                ofi_operaciones = str(i[0])+". "+str(i[1])
                
                
        query = "SELECT * FROM PERSONAL WHERE CARGO = 'S3'"
        efectivos_toe_SUB = b_d_i.INFORMACION_PEL(self, query)

        for i in efectivos_toe_SUB:

                sub_operaciones = str(i[0])+". "+str(i[1])

        table3 = doc.add_table(rows=1, cols=15) 
        table3.style = 'firmas' 
        row = table3.rows[0].cells 

        row[4].merge(row[0])
        row[0].paragraphs[0].add_run(sub_operaciones).bold = True
        row[9].merge(row[5])
        row[5].paragraphs[0].add_run(ofi_operaciones).bold = True
        row[14].merge(row[10])
        row[10].paragraphs[0].add_run(cdte).bold = True

        row = table3.add_row().cells 
        row[4].merge(row[0])
        row[0].paragraphs[0].add_run("Suboficial de Operaciones " + batallon).bold = False
        # row[0].text = str()
        row[9].merge(row[5])
        row[5].paragraphs[0].add_run("Oficial de Operaciones  "+ batallon).bold = False
        # row[5].text = str("Oficial de Operaciones  "+ linea_unidad)
        row[14].merge(row[10])
        row[10].paragraphs[0].add_run("Comandante " + batallon).bold = False
        # row[10].text = str("Comandante " + linea_unidad)


        directorio = APP_PATH+"/../datos/WORD"

        doc.save(directorio+"/" + nombre +".docx")

    except:
        print("inicio")

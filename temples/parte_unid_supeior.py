from docxtpl import DocxTemplate
from datetime import datetime
from base_datos.base_datos import *
#from mapa.distancias import *
import os
APP_PATH = os.getcwd()

concion_bd = Base_datos_central
def query_bd(self):
    def comando_query(self, **kwargs):
        datos = concion_bd.read_query_dinamico(self, **kwargs)
        return datos
    return  comando_query

#Llenado de los datos del primer formulario
@query_bd
def datos_base_datos(self, **kwargs):
    pass

directorio = APP_PATH +"/../datos/base_datos_no_borrar/temples/plantilla_maniobra.docx"
doc = DocxTemplate(directorio)
def suma(i):
      
            TC =0
            MY = 0
            CT = 0
            TE = 0
            ST = 0
            SP = 0
            SV = 0
            SS = 0
            CP = 0
            CS = 0
            C3 = 0

            if i[52]=="TC":
                    TC = TC +1 
            elif i[52]=="MY":
                    MY = MY +1
            elif i[52]=="CT":
                    CT = CT +1
            elif i[52]=="TE":
                    TE = TE +1
            elif i[52]=="ST":
                    ST = ST +1
            elif i[52]=="SP":
                    SP = SP +1 
            elif i[52]=="SV":
                    SV = SV +1
            elif i[52]=="SS":
                    SS = SS +1
            elif i[52]=="CP":
                    CP = CP +1
            elif i[52]=="CS":
                    CS = CS +1 
            elif i[52]=="C3":
                    C3 = C3 +1 
            sumas = TC + MY + CT + TE + ST + SP + SV + SS + CP + CS + C3
            return[sumas, TC , MY , CT , TE , ST , SP , SV , SS , CP , CS , C3]



def crear_word_parte_personal_unidad(self):
                    
    fecha_informe = self.ui.fecha_insitop_2.text() 
    fecha=datetime.strptime(fecha_informe, '%d/%m/%Y')
    fecha_name = fecha.strftime('%Y-%m-%d ')

    division =  self.ui.division.currentText()
    brigada =  self.ui.brigada.currentText()
    sigla_batallon_5 =  self.ui.sigla_batallon_5.currentText()

    unidad = "EJC"
    
    if sigla_batallon_5 != "" and sigla_batallon_5 != "---":
        informacion ="select * from ESTADISTICA where fecha_elaboracion  = '{}' and unidad = '{}'".format(fecha_name, sigla_batallon_5)
        unidad = sigla_batallon_5
    elif brigada != "" and brigada != "---":
        informacion ="select * from ESTADISTICA where fecha_elaboracion  = '{}' and brigada = '{}'".format(fecha_name, brigada)
        unidad = brigada
    elif division != "" and division != "---":
        informacion ="select * from ESTADISTICA where fecha_elaboracion  = '{}' and division = '{}'".format(fecha_name, division)
        unidad = division
    else:
        informacion ="select * from ESTADISTICA where fecha_elaboracion  = '{}' ".format(fecha_name)

    estadistica = b_d_e.INFORMACION_PEL(self, informacion)

    nombre = str("parte") +" - " +str(unidad) + " - " + str(fecha_name)

    operaciones_ofensivas = 0
    operaciones_defensivas = 0
    operaciones_estabilidad = 0
    operaciones_adac = 0
    operaciones_no_aplica = 0

    pel_descanso = 0
    pel_entrenamiento = 0
    pel_area = 0
    total_pelotones=0
    operaciones_activas = 0

    pel_bombona = 0
    pel_tenerife = 0
    pel_junin = 0
    pel_mutis = 0
    pel_amazona = 0
    pel_paya = 0
    pel_barbula = 0
    pel_boyaca = 0
    pel_pantano = 0
    pel_gameza = 0
    pel_tarapaca = 0
    pel_cienega = 0
    total_loo= 0

    tolarl_loe=0

    cdt_ofi = 0
    cdt_sub = 0
    cdt_h = 0
    cdt_m = 0

    oficial_efectivos = 0
    suboficial_efectivos = 0
    soldados_slp_efectivos = 0
    soldados_sl18_efectivos = 0
    soldados_sl12_efectivos = 0
    soldados_h_efectivos = 0
    soldados_m_efectivos = 0

    oficial_area = 0
    suboficial_area = 0
    soldados_slp_area = 0
    soldados_sl18_area = 0
    soldados_sl12_area = 0
    soldados_h_area = 0
    soldados_m_area = 0

    oficial_entrenamiento = 0
    suboficial_entrenamiento = 0
    soldados_slp_entrenamiento = 0
    soldados_sl18_entrenamiento = 0
    soldados_sl12_entrenamiento = 0
    soldados_h_entrenamiento = 0
    soldados_m_entrenamiento = 0

    oficial_descanso = 0
    suboficial_descanso = 0
    soldados_slp_descanso = 0
    soldados_sl18_descanso = 0
    soldados_sl12_descanso = 0
    soldados_h_descanso = 0
    soldados_m_descanso = 0

    oficial_pdmad = 0
    suboficial_pdmad = 0
    soldados_slp_pdmad = 0
    soldados_sl18_pdmad = 0
    soldados_sl12_pdmad = 0
    soldados_h_pdmad = 0
    soldados_m_pdmad = 0

    oficial_pdmat = 0
    suboficial_pdmat = 0
    soldados_slp_pdmat = 0
    soldados_sl18_pdmat = 0
    soldados_sl12_pdmat = 0
    soldados_h_pdmat = 0
    soldados_m_pdmat = 0

    oficial_novedad = 0
    suboficial_novedad = 0
    soldados_slp_novedad = 0
    soldados_sl18_novedad = 0
    soldados_sl12_novedad = 0
    soldados_h_novedad = 0
    soldados_m_novedad = 0

    oficial_nov = 0
    suboficial_nov = 0
    soldados_slp_nov = 0
    soldados_sl18_nov = 0
    soldados_sl12_nov = 0
    soldados_h_nov = 0
    soldados_m_nov = 0

    total_pel_operaciones= 0

    for i in estadistica:
        if i[5] != "":
            total_pelotones = total_pelotones + int(i[5])
        if i[6] != "":
            pel_area = pel_area + int(i[6])
        if i[7] != "":
            pel_entrenamiento = pel_entrenamiento + int(i[7])
        if i[8] != "":
            pel_descanso = pel_descanso + int(i[8])

        #operaciones
        if i[73] != "":
            operaciones_ofensivas = operaciones_ofensivas + int(i[73])
        if i[74] != "":
            operaciones_defensivas = operaciones_defensivas + int(i[74])
        if i[75] != "":
            operaciones_estabilidad = operaciones_estabilidad + int(i[75])
        if i[76] != "":
            operaciones_adac = operaciones_adac + int(i[76])
        if i[77] != "":
            operaciones_no_aplica = operaciones_no_aplica + int(i[77])
        if i[79] != "":
            operaciones_activas = operaciones_activas + int(i[79])

        total_pel_operaciones = operaciones_ofensivas + operaciones_defensivas + operaciones_estabilidad + operaciones_adac + operaciones_no_aplica
        #loe y loo
            
        if i[80] != "":
            pel_bombona = pel_bombona + int(i[80])       
        if i[81] != "":
            pel_mutis = pel_mutis + int(i[81])     
        if i[82] != "":
            pel_paya = pel_paya + int(i[82])        
        if i[83] != "":
            pel_boyaca = pel_boyaca + int(i[83])          
        if i[84] != "":
            pel_pantano = pel_pantano + int(i[84])              
        if i[85] != "":
            pel_tarapaca = pel_tarapaca + int(i[85])
        if i[86] != "":
            total_loo = total_loo + int(i[86])

                    #loe y loo
            
        if i[87] != "":
            pel_tenerife = pel_tenerife + int(i[87])       
        if i[88] != "":
            pel_junin = pel_junin + int(i[88])     
        if i[89] != "":
            pel_amazona = pel_amazona + int(i[89])        
        if i[90] != "":
            pel_barbula = pel_barbula + int(i[90])          
        if i[91] != "":
            pel_gameza = pel_gameza + int(i[91])              
        if i[92] != "":
            pel_cienega = pel_cienega + int(i[92])
        if i[93] != "":
            tolarl_loe = tolarl_loe + int(i[93])

        if i[69] != "":
            cdt_ofi = cdt_ofi + int(i[69])
        if i[70] != "":
            cdt_sub = cdt_sub + int(i[70])
        if i[71] != "":
            cdt_h = cdt_h + int(i[71])
        if i[72] != "":
            cdt_m = cdt_m + int(i[72])

            
        if i[13] != "":
            oficial_efectivos = oficial_efectivos + int(i[13])
        if i[14] != "":
            suboficial_efectivos = suboficial_efectivos + int(i[14])
        if i[15] != "":
            soldados_slp_efectivos = soldados_slp_efectivos + int(i[15])
        if i[16] != "":
            soldados_sl18_efectivos = soldados_sl18_efectivos + int(i[16])
        if i[17] != "":
            soldados_sl12_efectivos = soldados_sl12_efectivos + int(i[17])
        if i[18] != "":
            soldados_h_efectivos = soldados_h_efectivos + int(i[18])
        if i[19] != "":
            soldados_m_efectivos = soldados_m_efectivos + int(i[19])

                        
        if i[20] != "":
            oficial_area = oficial_area + int(i[20])
        if i[21] != "":
            suboficial_area = suboficial_area + int(i[21])
        if i[22] != "":
            soldados_slp_area = soldados_slp_area + int(i[22])
        if i[23] != "":
            soldados_sl18_area = soldados_sl18_area + int(i[23])
        if i[24] != "":
            soldados_sl12_area = soldados_sl12_area + int(i[24])
        if i[25] != "":
            soldados_h_area = soldados_h_area + int(i[25])
        if i[26] != "":
            soldados_m_area = soldados_m_area + int(i[26])

    
        if i[27] != "":
            oficial_entrenamiento = oficial_entrenamiento + int(i[27])
        if i[28] != "":
            suboficial_entrenamiento = suboficial_entrenamiento + int(i[28])
        if i[29] != "":
            soldados_slp_entrenamiento = soldados_slp_entrenamiento + int(i[29])
        if i[30] != "":
            soldados_sl18_entrenamiento = soldados_sl18_entrenamiento + int(i[30])
        if i[31] != "":
            soldados_sl12_entrenamiento = soldados_sl12_entrenamiento + int(i[31])
        if i[32] != "":
            soldados_h_entrenamiento = soldados_h_entrenamiento + int(i[32])
        if i[33] != "":
            soldados_m_entrenamiento = soldados_m_entrenamiento + int(i[33])

                        
        if i[34] != "":
            oficial_descanso = oficial_descanso + int(i[34])
        if i[35] != "":
            suboficial_descanso = suboficial_descanso + int(i[35])
        if i[36] != "":
            soldados_slp_descanso = soldados_slp_descanso + int(i[36])
        if i[37] != "":
            soldados_sl18_descanso = soldados_sl18_descanso + int(i[37])
        if i[38] != "":
            soldados_sl12_descanso = soldados_sl12_descanso + int(i[38])
        if i[39] != "":
            soldados_h_descanso = soldados_h_descanso + int(i[39])
        if i[40] != "":
            soldados_m_descanso = soldados_m_descanso + int(i[40])

                
        if i[41] != "":
            oficial_pdmad = oficial_pdmad + int(i[41])
        if i[42] != "":
            suboficial_pdmad = suboficial_pdmad + int(i[42])
        if i[43] != "":
            soldados_slp_pdmad = soldados_slp_pdmad + int(i[43])
        if i[44] != "":
            soldados_sl18_pdmad = soldados_sl18_pdmad + int(i[44])
        if i[45] != "":
            soldados_sl12_pdmad = soldados_sl12_pdmad + int(i[45])
        if i[46] != "":
            soldados_h_pdmad = soldados_h_pdmad + int(i[46])
        if i[47] != "":
            soldados_m_pdmad = soldados_m_pdmad + int(i[47])

    
        if i[48] != "":
            oficial_pdmat = oficial_pdmat + int(i[48])
        if i[49] != "":
            suboficial_pdmat = suboficial_pdmat + int(i[49])
        if i[50] != "":
            soldados_slp_pdmat = soldados_slp_pdmat + int(i[50])
        if i[51] != "":
            soldados_sl18_pdmat = soldados_sl18_pdmat + int(i[51])
        if i[52] != "":
            soldados_sl12_pdmat = soldados_sl12_pdmat + int(i[52])
        if i[53] != "":
            soldados_h_pdmat = soldados_h_pdmat + int(i[53])
        if i[54] != "":
            soldados_m_pdmat = soldados_m_pdmat + int(i[54])

                
        if i[55] != "":
            oficial_novedad = oficial_novedad + int(i[55])
        if i[56] != "":
            suboficial_novedad = suboficial_novedad + int(i[56])
        if i[57] != "":
            soldados_slp_novedad = soldados_slp_novedad + int(i[57])
        if i[58] != "":
            soldados_sl18_novedad = soldados_sl18_novedad + int(i[58])
        if i[59] != "":
            soldados_sl12_novedad = soldados_sl12_novedad + int(i[59])
        if i[60] != "":
            soldados_h_novedad = soldados_h_novedad + int(i[60])
        if i[61] != "":
            soldados_m_novedad = soldados_m_novedad + int(i[61])

   
        if i[62] != "":
            oficial_nov = oficial_nov + int(i[62])
        if i[63] != "":
            suboficial_nov = suboficial_nov + int(i[63])
        if i[64] != "":
            soldados_slp_nov = soldados_slp_nov + int(i[64])
        if i[65] != "":
            soldados_sl18_nov = soldados_sl18_nov + int(i[65])
        if i[66] != "":
            soldados_sl12_nov = soldados_sl12_nov + int(i[66])
        if i[67] != "":
            soldados_h_nov = soldados_h_nov + int(i[67])
        if i[68] != "":
            soldados_m_nov = soldados_m_nov + int(i[68])



    lambda_func = lambda x: x if x != 0 else ""

        #llenado de parte de personal en el informe de INSITOP


    context = { 

        'fecha' :  fecha_name,

        'UNIDAD': unidad,

        'total_pelotones': lambda_func(total_pelotones),
        'pel_area' : lambda_func(pel_area),
        'pel_entrenamiento' : lambda_func(pel_entrenamiento),
        'pel_descanso' : lambda_func(pel_descanso),

        'cdte_ofi' : lambda_func(cdt_ofi),
        'cdte_sub' : lambda_func(cdt_sub),
        'cdte_h' : lambda_func(cdt_h),
        'cdte_m' : lambda_func(cdt_m),

        'efectivos_ofi' : lambda_func(oficial_efectivos),
        'efectivos_sub' : lambda_func(suboficial_efectivos),
        'efectivos_slp' : lambda_func(soldados_slp_efectivos),
        'efectivos_sl18' : lambda_func(soldados_sl18_efectivos),
        'efectivos_sl12' : lambda_func(soldados_sl12_efectivos),
        'efectivos_h' : lambda_func(soldados_h_efectivos),
        'efectivos_m' : lambda_func(soldados_m_efectivos),
 
        'area_op_ofi' : lambda_func(oficial_area),
        'area_op_sub' : lambda_func(suboficial_area),
        'area_op_slp' : lambda_func(soldados_slp_area),
        'area_op_sl18' : lambda_func(soldados_sl18_area),
        'area_op_sl12' : lambda_func(soldados_sl12_area),
        'area_op_h' : lambda_func(soldados_h_area),
        'area_op_m' : lambda_func(soldados_m_area),

        'entrenamiento_ofi' : lambda_func(oficial_entrenamiento),
        'entrenamiento_sub' : lambda_func(suboficial_entrenamiento),
        'entrenamiento_slp' : lambda_func(soldados_slp_entrenamiento),
        'entrenamiento_sl18' : lambda_func(soldados_sl18_entrenamiento),
        'entrenamiento_sl12' : lambda_func(soldados_sl12_entrenamiento),
        'entrenamiento_h' : lambda_func(soldados_h_entrenamiento),
        'entrenamiento_m' : lambda_func(soldados_m_entrenamiento),

        'descando_ofi' : lambda_func(oficial_descanso),
        'descando_sub' : lambda_func(suboficial_descanso),
        'descando_slp' : lambda_func(soldados_slp_descanso),
        'descando_sl18' : lambda_func(soldados_sl18_descanso),
        'descando_sl12' : lambda_func(soldados_sl12_descanso),
        'descando_h' : lambda_func(soldados_h_descanso),
        'descando_m' : lambda_func(soldados_m_descanso),

        'pdmad_ofi' : lambda_func(oficial_pdmad),
        'pdmad_sub' : lambda_func(suboficial_pdmad),
        'pdmad_slp' : lambda_func(soldados_slp_pdmad),
        'pdmad_sl18' : lambda_func(soldados_sl18_pdmad),
        'pdmad_sl12' : lambda_func(soldados_sl12_pdmad),
        'pdmad_h' : lambda_func(soldados_h_pdmad),
        'pdmad_m' : lambda_func(soldados_m_pdmad),

        'pmdat_ofi' : lambda_func(oficial_pdmat),
        'pmdat_sub' : lambda_func(suboficial_pdmat),
        'pmdat_slp' : lambda_func(soldados_slp_pdmat),
        'pmdat_sl18' : lambda_func(soldados_sl18_pdmat),
        'pmdat_sl12' : lambda_func(soldados_sl12_pdmat),
        'pmdat_h' : lambda_func(soldados_h_pdmat),
        'pmdat_m' : lambda_func(soldados_m_pdmat),

        'nov_recuperables_ofi' : lambda_func(oficial_novedad),
        'nov_recuperables_sub' : lambda_func(suboficial_novedad),
        'nov_recuperables_slp' : lambda_func(soldados_slp_novedad),
        'nov_recuperables_sl18' : lambda_func(soldados_sl18_novedad),
        'nov_recuperables_sl12' : lambda_func(soldados_sl12_novedad),
        'nov_recuperables_h' : lambda_func(soldados_h_novedad),
        'nov_recuperables_m' : lambda_func(soldados_m_novedad),

        'nov_no_recuperables_ofi' : lambda_func(oficial_nov),
        'nov_no_recuperables_sub' : lambda_func(suboficial_nov),
        'nov_no_recuperables_slp' : lambda_func(soldados_slp_nov),
        'nov_no_recuperables_sl18' : lambda_func(soldados_sl18_nov),
        'nov_no_recuperables_sl12' : lambda_func(soldados_sl12_nov),
        'nov_no_recuperables_h' : lambda_func(soldados_h_nov),
        'nov_no_recuperables_m' : lambda_func(soldados_m_nov),


        'pel_op_ofensiva' : lambda_func(operaciones_ofensivas),
        'pel_op_defensiva' : lambda_func(operaciones_defensivas),
        'pel_op_estabilidad' : lambda_func(operaciones_estabilidad),
        'pel_op_adac' : lambda_func(operaciones_adac),
        'pel_no_aplica' : lambda_func(operaciones_no_aplica),
        'total_op' : lambda_func(total_pel_operaciones),
        'op_activas' : lambda_func(operaciones_activas),

        'bobona' : lambda_func(pel_bombona),
        'mutis' : lambda_func(pel_mutis),
        'paya' : lambda_func(pel_paya),
        'boyaca' : lambda_func(pel_boyaca),
        'pantano_pargas' : lambda_func(pel_pantano),
        'tarapaca' : lambda_func(pel_tarapaca),
        'total_loo' : lambda_func(total_loo),

        'tenerife' : lambda_func(pel_tenerife), 
        'junin' : lambda_func(pel_junin), 
        'amazonia' : lambda_func(pel_amazona), 
        'barbula' : lambda_func(pel_barbula), 
        'gameza' : lambda_func(pel_gameza), 
        'cienega' : lambda_func(pel_cienega), 
        'tolarl_loe' : lambda_func(tolarl_loe)

                
    }


    doc.render(context)

    import docx 
    
    doc2 = docx.Document() 
    
    doc2.add_heading('GeeksForGeeks', 0) 
    


    
    # table = doc.add_table(rows=1, cols=2) 
    # table.style = 'resultados' 
    # row = table.rows[0].cells 
    # row[0].text = 'Id'
    # row[1].text = 'Name'
    
    # for id, name in tabla_1: 
    
        
    #     row = table.add_row().cells 
        
    #     row[0].text = str(id) 
    #     row[1].text = name 

    # # Añadimos un párrafo
    # parrafo = doc.add_paragraph()
    # # parrafo.add_run().add_break()

    directorio = APP_PATH+"/../datos/WORD"

    doc.save(directorio+"/" + nombre +".docx")
